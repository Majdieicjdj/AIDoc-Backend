from io import BytesIO
from flask import Flask, request, jsonify, send_from_directory
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from datetime import datetime
import re, json
from flask_cors import CORS
from dotenv import load_dotenv
from groq import Groq
from duckduckgo_search import DDGS
import requests
import uuid
from typing import List, Dict, Optional,Any
from dataclasses import dataclass, field
from enum import Enum, auto
load_dotenv()

app = Flask(__name__)
CORS(app, expose_headers=['X-Filename'])

# Configure Google API
genai.configure(api_key=os.environ['GEMINI_API_KEY'])

# Initialize Gemini model
model = genai.GenerativeModel('gemini-1.5-flash')
client = Groq(max_retries=3)
ddgs = DDGS()

def create_element(name):
    """
    Create an OxmlElement with the given name
    """
    return OxmlElement(name)

def create_attribute(element, name, value):
    """
    Create an attribute for the given element
    """
    element.set(qn(name), value)
    
class AgentStatus(Enum):
    ACTIVE = auto()
    INACTIVE = auto()
    TEMPLATE = auto()

@dataclass
class Agent:
    id: str
    name: str
    icon: str
    status: AgentStatus = AgentStatus.ACTIVE
    configuration: Dict[str, Any] = field(default_factory=dict)

class AgentManager:
    def __init__(self):
        self._agents: List[Agent] = []
        self._initialize_default_agents()

    def _initialize_default_agents(self):
        """Initialize default agents with basic configurations."""
        default_agents = [
            Agent(
                id="template_agent", 
                name="Template Agent", 
                icon="🔖", 
                status=AgentStatus.TEMPLATE
            ),
            # Add more default agents as needed
        ]
        self._agents.extend(default_agents)

    def get_agents(self) -> List[Agent]:
        """Retrieve all active and template agents."""
        return [agent for agent in self._agents if agent.status in [AgentStatus.ACTIVE, AgentStatus.TEMPLATE]]

    def add_agent(self, agent: Agent) -> None:
        """Add a new agent to the system."""
        if any(existing.id == agent.id for existing in self._agents):
            raise ValueError(f"Agent with ID {agent.id} already exists")
        self._agents.append(agent)

    def remove_agent(self, agent_id: str) -> None:
        """Remove an agent by its ID."""
        self._agents = [agent for agent in self._agents if agent.id != agent_id]

    def get_agent_by_id(self, agent_id: str) -> Optional[Agent]:
        """Retrieve a specific agent by its ID."""
        return next((agent for agent in self._agents if agent.id == agent_id), None)

    def activate_template_agent(self, template_id: str, configuration: Dict[str, Any]) -> Agent:
        """Activate a template agent with specific configuration."""
        template_agent = self.get_agent_by_id(template_id)
        if not template_agent or template_agent.status != AgentStatus.TEMPLATE:
            raise ValueError("Invalid template agent")

        new_agent = Agent(
            id=f"agent_{len(self._agents) + 1}",
            name=f"{template_agent.name} Instance",
            icon=template_agent.icon,
            status=AgentStatus.ACTIVE,
            configuration=configuration
        )
        self.add_agent(new_agent)
        return new_agent
    
class AgentTemplateType(Enum):
    CUSTOM = auto()
    SYSTEM = auto()

# Now use AgentTemplateType in the AgentTemplate class
@dataclass
class AgentTemplate:
    id: str
    name: str
    description: Optional[str] = None
    type: AgentTemplateType = AgentTemplateType.CUSTOM
    file_path: Optional[str] = None
    created_at: datetime = field(default_factory=datetime.now)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
class AgentTemplateManager:
        def __init__(self, base_template_dir: str = 'agent_templates'):
            """
            Initialize template management system with a base directory for storing templates
            
            Args:
                base_template_dir (str): Directory to store uploaded agent templates
            """
            self.base_template_dir = base_template_dir
            os.makedirs(base_template_dir, exist_ok=True)
            self._templates: List[AgentTemplate] = []

        def upload_template(self, file_path: str, template_name: Optional[str] = None) -> AgentTemplate:
            """
            Upload and process an agent template file
            
            Args:
                file_path (str): Path to the template file
                template_name (Optional[str]): Custom name for the template
            
            Returns:
                AgentTemplate: Created template object
            """
            # Validate file
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Template file not found: {file_path}")
            
            # Generate unique ID and filename
            template_id = str(uuid.uuid4())
            file_extension = os.path.splitext(file_path)[1]
            new_filename = f"{template_id}{file_extension}"
            
            # Determine template name
            if template_name is None:
                template_name = os.path.splitext(os.path.basename(file_path))[0]
            
            # Copy file to templates directory
            destination_path = os.path.join(self.base_template_dir, new_filename)
            os.rename(file_path, destination_path)
            
            # Create template object
            template = AgentTemplate(
                id=template_id,
                name=template_name,
                file_path=destination_path,
                type=AgentTemplateType.CUSTOM,
                metadata={
                    'original_filename': os.path.basename(file_path),
                    'file_size': os.path.getsize(destination_path)
                }
            )
            
            self._templates.append(template)
            return template

        def list_templates(self) -> List[AgentTemplate]:
            """
            List all available templates
            
            Returns:
                List[AgentTemplate]: Available templates
            """
            return self._templates

        def get_template_by_id(self, template_id: str) -> Optional[AgentTemplate]:
            """
            Retrieve a specific template by ID
            
            Args:
                template_id (str): Unique template identifier
            
            Returns:
                Optional[AgentTemplate]: Template if found, None otherwise
            """
            return next((template for template in self._templates if template.id == template_id), None)

        def delete_template(self, template_id: str) -> bool:
            """
            Delete a template by its ID
            
            Args:
                template_id (str): Unique template identifier
            
            Returns:
                bool: Whether deletion was successful
            """
            template = self.get_template_by_id(template_id)
            if template:
                # Remove file
                if os.path.exists(template.file_path):
                    os.remove(template.file_path)
                
                # Remove from templates list
                self._templates = [t for t in self._templates if t.id != template_id]
                return True
            return False

        def export_template_catalog(self, output_path: str = 'template_catalog.json'):
            """
            Export template catalog as JSON
            
            Args:
                output_path (str): Path to save template catalog
            """
            catalog = [{
                'id': template.id,
                'name': template.name,
                'type': template.type.name,
                'created_at': template.created_at.isoformat(),
                'metadata': template.metadata
            } for template in self._templates]
            
            with open(output_path, 'w') as f:
                json.dump(catalog, f, indent=2)

# Example usage
template_manager = AgentTemplateManager()

def handle_template_upload(file_path: str):
    """
    Handler for template file upload
    
    Args:
        file_path (str): Path to uploaded template file
    """
    try:
        template = template_manager.upload_template(file_path)
        print(f"Template uploaded: {template.name} (ID: {template.id})")
    except Exception as e:
        print(f"Template upload failed: {e}")

# Simulated file upload handler (would typically be in a web framework like Flask/FastAPI)
def upload_template_file(uploaded_file):
    """
    Process uploaded file from web interface
    
    Args:
        uploaded_file: File object from web framework
    """
    # Save temporary file
    temp_path = os.path.join('temp_uploads', uploaded_file.filename)
    uploaded_file.save(temp_path)
    
    # Process template
    handle_template_upload(temp_path)
        
        
        


# Example usage
agent_manager = AgentManager()

def handle_template_click():
    """
    Example handler for template agent activation.
    In a real application, this would likely involve user interaction 
    to provide configuration details.
    """
    template_agent = agent_manager.get_agent_by_id("template_agent")
    if template_agent:
        try:
            new_agent = agent_manager.activate_template_agent(
                "template_agent", 
                configuration={"created_at": datetime.now()}
            )
            print(f"Created new agent: {new_agent.name}")
        except ValueError as e:
            print(f"Error creating agent: {e}")


def add_page_number(paragraph):
    """
    Add page number field to the given paragraph
    """
    # Create run for "Page "
    page_run = paragraph.add_run()
    page_run.font.size = Pt(10)
    page_run.font.name = 'Arial'
    page_run.text = "Page "

    # Create and add page number field
    fld_char1 = create_element('w:fldChar')
    create_attribute(fld_char1, 'w:fldCharType', 'begin')

    instr_text = create_element('w:instrText')
    create_attribute(instr_text, 'xml:space', 'preserve')
    instr_text.text = "PAGE"

    fld_char2 = create_element('w:fldChar')
    create_attribute(fld_char2, 'w:fldCharType', 'end')

    run = paragraph.add_run()
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)

    # Add " of " text
    of_run = paragraph.add_run()
    of_run.font.size = Pt(10)
    of_run.font.name = 'Arial'
    of_run.text = " of "

    # Create and add total pages field
    fld_char3 = create_element('w:fldChar')
    create_attribute(fld_char3, 'w:fldCharType', 'begin')

    instr_text2 = create_element('w:instrText')
    create_attribute(instr_text2, 'xml:space', 'preserve')
    instr_text2.text = "NUMPAGES"

    fld_char4 = create_element('w:fldChar')
    create_attribute(fld_char4, 'w:fldCharType', 'end')

    run = paragraph.add_run()
    run._r.append(fld_char3)
    run._r.append(instr_text2)
    run._r.append(fld_char4)

def add_formatted_paragraph(doc, text, font_name='Arial', font_size=11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """
    Add a formatted paragraph to the document
    """
    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment
    
    # Split text into paragraphs if multiple are provided
    paragraphs = text.split('\n\n')
    
    for i, para_text in enumerate(paragraphs):
        if i > 0:
            # Add a new paragraph for subsequent paragraphs
            paragraph = doc.add_paragraph()
            paragraph.alignment = alignment
            
        run = paragraph.add_run(para_text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
    
    return paragraph

def create_document_with_footer(content, doc_id, image_urls=None):
    """
    Create a document with content, images after each section, and page numbers in the footer
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading(content['title'], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content sections
    for idx, section in enumerate(content['sections']):
        # Add section heading
        section_heading = doc.add_heading(section['heading'], level=2)
        section_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add section content with proper formatting
        add_formatted_paragraph(
            doc, 
            section['content'],
            font_name='Arial',
            font_size=11,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        )
        
        # Add image after each section (if available)
        if image_urls and idx < len(image_urls):
            image_url = image_urls[idx]
            try:
                # Download image
                response = requests.get(image_url)
                img_stream = BytesIO(response.content)
                # Add image to the document
                image_paragraph = doc.add_paragraph()
                run = image_paragraph.add_run()
                run.add_picture(img_stream, width=Inches(5))  # Adjust width as necessary
                image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER            
            except Exception as e:
                print(f"Error downloading image from {image_url}: {e}")
        
        # Add spacing after section
        doc.add_paragraph()
    
    # Add page numbers in footer
    sections = doc.sections
    for section in sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(footer_para)
    
    # Create directory if it doesn't exist
    os.makedirs('generated_docs', exist_ok=True)
    
    # Save document
    filename = f"generated_docs/{doc_id}.docx"
    doc.save(filename)
    return filename

def extract_json(text):
    """
    Extract JSON content from a given string, even if it's not enclosed in code identifiers.
    """
    # Regex to match JSON objects by detecting curly braces
    json_pattern = r'\{(?:[^{}]*|\{.*?\})*\}'
    
    # Search for JSON-like patterns in the text
    matches = re.findall(json_pattern, text, re.DOTALL)
    
    for match in matches:
        try:
            # Attempt to parse the match as JSON
            return json.loads(match)
        except json.JSONDecodeError:
            continue
    
    # Raise an error if no valid JSON is found
    raise ValueError("No valid JSON content found in the input text.")

def generate_content(prompt):
    """
    Generate content using Gemini API
    """
    system_prompt = """
    Generate a document with the following structure:
    1. A title related to the topic
    2. Multiple sections, each with:
       - A relevant heading
       - Detailed content (minimum 2 paragraphs separated by double newlines)
    The output must only be in the JSON object only, no other sentences or after the output response, only JSON object in the given format must be provided.
    Return the response as a JSON object with this structure:
    {
        "title": "Main Title",
        "sections": [
            {
                "heading": "Section Heading",
                "content": "First paragraph content.\n\nSecond paragraph content."
            }
        ]
    }
    Make sure paragraphs are separated by double newlines (\n\n) in the content.
    """
    
    try:
        # Combine system prompt and user prompt
        full_prompt = f"{system_prompt}\n\nTopic: {prompt}"
        
        # images = ddgs.images(prompt, max_results=10)
        image_urls = []
        # for image in images:
        #     image_urls.append(image['image'])
        #     print(image['image'])
        # print('\n\n')
        
        # while len(image_urls) > 4:
        #     image_urls.pop()
        
        # Generate content
        # response = model.generate_content(full_prompt)
        # content = extract_json(response.text)
        
        completion = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {
                    "role": "system",
                    "content": "Generate a detailed document with the following structure:\n    1. A title related to the topic\n    2. Multiple sections, each with:\n       - A relevant heading\n       - Detailed content (minimum 2 paragraphs separated by double newlines)\n    The output must only be in the JSON object only, no other sentences or after the output response, only JSON object in the given format must be provided.\n    Return the response as a JSON object with this structure:\n    {\n        \"title\": \"Main Title\",\n        \"sections\": [\n            {\n                \"heading\": \"Section Heading\",\n                \"content\": \"First paragraph content.\\n\\nSecond paragraph content.\"\n            }\n        ]\n    }\n    Make sure paragraphs are separated by double newlines (\\n\\n) in the content.\nThe output must be only in JSON object format."
                },
                {
                    "role": "user",
                    "content": full_prompt
                },
            ],
            temperature=0.5,
            max_tokens=7024,
            top_p=1,
            stream=False,
            response_format={"type": "json_object"},
            stop=None,
        )

        content = completion.choices[0].message.content
        content = extract_json(content)

        print(content)

        print('/n/nheadings---------------------------------\n')
        for section in content['sections']:
            print(section['heading'])
            images = ddgs.images(section['heading'], max_results=5)
            print(images)
            print('/n/n')
            
            if len(images):
                for image in images:
                    try:
                        response = requests.get(image['image'])
                        
                        if response.status_code == 200:
                            image_urls.append(image['image'])
                            break
                    except Exception as e:
                        print(e) 
                        continue
                    
                    print(image['image'])
                print('\n\n')
            
        print('/n/n--------------------------------------------')
            
        return content, image_urls
    
    except Exception as e:
        raise Exception(f"Error generating content: {str(e)}")

@app.route('/')
def status():
    return 'Server is running'

@app.route('/generate-document', methods=['POST'])
def generate_document():
    try:
        # Get prompt and document ID from request
        data = request.get_json()
        if not data or 'prompt' not in data or 'doc_id' not in data:
            return jsonify({
                'error': 'Missing required fields: prompt and doc_id'
            }), 400
            
        prompt = data['prompt']
        doc_id = data['doc_id']
        
        print('\n\n\n----------------------\nPrompt: ', prompt, '\n\n====================================')
        
        # Generate content using Gemini
        content, image_urls = generate_content(prompt)
        
        # Create document with the generated content
        filename = create_document_with_footer(content, doc_id, image_urls)
        
        # return jsonify({
        #     'status': 'success',
        #     'message': 'Document generated successfully',
        #     'filename': filename,
        #     'generated_at': datetime.now().isoformat()
        # }), 200
        print('\n\nFilename: ')
        filename = filename.split('/')[-1]
        print(f'filename: {filename}')
        print(filename)
        response = send_from_directory(
            directory='./generated_docs',
            path=filename,
            as_attachment=True
        )
        response.headers['X-Filename'] = filename  # Add custom header
        print(response.headers)
        return response
        
    except Exception as e:
        return jsonify({
            'error': str(e)
        }), 500

# Error handlers
@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error'}), 500

if __name__ == '__main__':
    app.run(debug=True)